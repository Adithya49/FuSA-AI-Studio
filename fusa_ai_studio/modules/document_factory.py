from __future__ import annotations

import ast
import json
from pathlib import Path

import streamlit as st

from fusa_ai_studio.core.services import Services
from fusa_ai_studio.database.repository import now_iso
from fusa_ai_studio.ui.components import data_table
from fusa_ai_studio.ui.genai import render_ai_response_with_chat, run_genai_action
from fusa_ai_studio.ui.workproduct_inputs import render_workproduct_inputs


def _try_parse_object(text: str):
    text = text.strip()
    if not text:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    try:
        return ast.literal_eval(text)
    except (ValueError, SyntaxError):
        return None


def _render_object_table(doc, obj) -> bool:
    if isinstance(obj, dict):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Field"
        header_cells[1].text = "Value"
        for key, value in obj.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
        return True

    if isinstance(obj, list) and obj and all(isinstance(item, dict) for item in obj):
        keys = sorted({key for item in obj for key in item.keys()})
        if not keys:
            return False
        table = doc.add_table(rows=1, cols=len(keys))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, key in enumerate(keys):
            header_cells[index].text = str(key)
        for item in obj:
            row_cells = table.add_row().cells
            for index, key in enumerate(keys):
                row_cells[index].text = str(item.get(key, ""))
        return True

    return False


def _write_docx(path: Path, title: str, content: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    lines = content.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            index += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            index += 1
            continue

        if line.startswith("- "):
            remainder = line[2:].strip()
            parsed = _try_parse_object(remainder)
            if parsed is not None:
                table_rows = [parsed] if isinstance(parsed, dict) else parsed
                next_index = index + 1
                while next_index < len(lines) and lines[next_index].startswith("- "):
                    next_parsed = _try_parse_object(lines[next_index][2:].strip())
                    if isinstance(next_parsed, dict) and isinstance(parsed, dict):
                        table_rows.append(next_parsed)
                        next_index += 1
                        continue
                    break
                if isinstance(parsed, dict):
                    if _render_object_table(doc, table_rows):
                        index = next_index
                        continue
                if _render_object_table(doc, parsed):
                    index += 1
                    continue
            doc.add_paragraph(line[2:], style="List Bullet")
            index += 1
            continue

        if line.strip().startswith("{") or line.strip().startswith("["):
            block_lines = []
            next_index = index
            while next_index < len(lines) and lines[next_index].strip() and not lines[next_index].startswith("# ") and not lines[next_index].startswith("- "):
                block_lines.append(lines[next_index])
                next_index += 1
            parsed = _try_parse_object("\n".join(block_lines))
            if parsed is not None and _render_object_table(doc, parsed):
                index = next_index
                continue

        parsed = _try_parse_object(line)
        if parsed is not None and _render_object_table(doc, parsed):
            index += 1
            continue

        if line.strip():
            doc.add_paragraph(line)
        index += 1
    doc.save(path)

def _artifact_summary(services: Services, project_id: str) -> str:
    repo = services.repo

    sections = [
        ("items", "Items"),
        ("hazards", "HARA"),
        ("safety_goals", "Safety Goals"),
        ("fsc_requirements", "FSC Requirements"),
        ("tsc_requirements", "TSC Requirements"),
    ]

    lines = []

    for table_name, title in sections:
        rows = repo.list_table(table_name, project_id)

        lines.append(f"## {title}")

        if not rows:
            lines.append("No records found.")
            lines.append("")
            continue

        for idx, row in enumerate(rows, start=1):
            lines.append(f"### Record {idx}")

            clean_row = {}

            for key, value in row.items():
                if value is None:
                    clean_row[key] = ""
                elif isinstance(value, (dict, list)):
                    clean_row[key] = json.dumps(value, indent=2)
                else:
                    clean_row[key] = str(value)

            lines.append(json.dumps(clean_row))
            lines.append("")

    return "\n".join(lines)


def _trace_summary(services: Services, project_id: str) -> str:
    links = services.repo.list_table("trace_links", project_id)
    return "\n".join(f"- {link['source_type']}:{link['source_id']} -> {link['target_type']}:{link['target_id']} ({link['link_type']})" for link in links)


def _write_docx(path: Path, title: str, content: str) -> None:
    from docx import Document

    doc = Document()

    doc.add_heading(title, level=1)

    lines = content.splitlines()

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue

        # JSON record -> single table
        if line.startswith("{") and line.endswith("}"):
            try:
                data = json.loads(line)

                if isinstance(data, dict):
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"

                    header = table.rows[0].cells
                    header[0].text = "Field"
                    header[1].text = "Value"

                    for key, value in data.items():
                        row = table.add_row().cells
                        row[0].text = str(key)

                        if value is None:
                            value = ""

                        row[1].text = str(value)

                    doc.add_paragraph()

                    continue

            except Exception:
                pass

        # Normal paragraph
        doc.add_paragraph(line)

    doc.save(path)

def render(services: Services, project_id: str) -> None:
    repo = services.repo
    st.title("Document Factory")
    input_tab, generation_tab, templates_tab, records_tab = st.tabs(["Inputs", "Generate", "Templates", "Records"])
    with input_tab:
        render_workproduct_inputs(services, project_id, "Document Factory")
    with generation_tab:
        project = repo.get_project(project_id) or {}
        doc_type = st.selectbox("Document", ["Item Definition Report", "HARA Report", "Safety Goals Report", "FSC Report", "TSC Report", "Traceability Matrix", "Safety Case Summary", "AI Gap Assessment"])
        output_format = st.selectbox("Format", ["Markdown", "DOCX", "TXT"])
        if st.button("Generate document", type="primary"):
            artifact_summary = _artifact_summary(services, project_id)
            trace_summary = _trace_summary(services, project_id)
            if doc_type == "AI Gap Assessment":
                answer = run_genai_action(
                    "Document Factory",
                    lambda: services.rag.ask(project_id, "Document Factory", "Generate an auditable safety case gap assessment with cited project context."),
                )
                body = answer.text
                render_ai_response_with_chat(services, project_id, "Document Factory", answer, "document_factory_ai")
            else:
                body = f"# {doc_type}\n\n## Project\n{project.get('name', project_id)}\n\n{project.get('description', '')}\n\n## Artifact Summary\n{artifact_summary}\n\n## Traceability Summary\n{trace_summary}\n"
            safe_name = doc_type.lower().replace(" ", "_")
            suffix = {"Markdown": "md", "DOCX": "docx", "TXT": "txt"}[output_format]
            path = services.config.export_dir / f"{safe_name}_{now_iso().replace(':', '').replace('Z', '')}.{suffix}"
            if output_format == "DOCX":
                _write_docx(path, doc_type, body)
            else:
                path.write_text(body, encoding="utf-8")
            repo.insert(
                "documents",
                {
                    "project_id": project_id,
                    "title": doc_type,
                    "doc_type": output_format,
                    "content": body,
                    "file_path": str(path),
                },
            )
            repo.add_memory(project_id, "document", f"Generated {doc_type} at {path.name}.", 2)
            st.success(f"Generated {path}")
            st.download_button("Download document", data=path.read_bytes(), file_name=path.name)
    with templates_tab:
        data_table(repo.list_table("doc_templates"), "No document templates recorded.")
    with records_tab:
        data_table(repo.list_table("documents", project_id), "No generated documents recorded.")
